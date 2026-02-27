# =====================================================
# Resume Classification 
# =====================================================

import streamlit as st
import re
import joblib
import pandas as pd
import pdfplumber
from datetime import datetime
from nltk.tokenize import RegexpTokenizer
import spacy
import nltk

from nltk.corpus import stopwords
from spacy.cli import download

# -----------------------------------------------------
# PAGE CONFIG
# -----------------------------------------------------

st.set_page_config(
    page_title="Resume Classification System",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Resume Classification System")
st.markdown("---")

# -----------------------------------------------------
# ENSURE NLP DEPENDENCIES
# -----------------------------------------------------

nltk.download("stopwords")

try:
    nlp = spacy.load("en_core_web_sm")
except:
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# -----------------------------------------------------
# LOAD MODEL
# -----------------------------------------------------

@st.cache_resource
def load_model():
    model = joblib.load("models/resume_classifier_nb_pipeline.pkl")
    label_encoder = joblib.load("models/label_encoder.pkl")
    return model, label_encoder

model, label_encoder = load_model()

# -----------------------------------------------------
# TEXT EXTRACTION
# -----------------------------------------------------

def extract_text_from_pdf(file):
    full_text = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
    return " ".join(full_text)

def extract_text_from_docx(file):
    from docx import Document
    doc = Document(file)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    return " ".join(full_text)

def extract_text(uploaded_file):
    ext = uploaded_file.name.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(uploaded_file)
    elif ext == "docx":
        return extract_text_from_docx(uploaded_file)
    else:
        st.error("Only PDF and DOCX files are supported.")
        return ""

# -----------------------------------------------------
# REGEX PATTERNS
# -----------------------------------------------------

tokenizer = RegexpTokenizer(r'[A-Za-z][A-Za-z0-9\+\#\.-]*')

# ----------------------------
# Regex patterns
# ----------------------------

TECH_PATTERN = re.compile(
    r"""
    (
        (html|css|v|ec|pt)  |
        c(\+\+|\#)        |
        [a-z]+[0-9]+      |
        [0-9]+[a-z]+      |
        [a-z]+[_-][a-z0-9]+ |
        [a-z]{3,}(?:[a-z0-9]+)?
    )
    """,
    re.X
)

FLUFF_PATTERN = re.compile(
    r"""
    ^(
        responsibility|responsibilities|
        professional|profile|summary|objective|
        organization|company|
        education|college|university|
        information|details|
        customer|client|user|
        service|services|
        work|experience|career|
        title|personal|dates|fields|also|additionally|various|active
    )s?$
    """,
    re.X
)

ACTION_VERB_PATTERN = re.compile(r'.*(ing|ed)$')

# ----------------------------
# Stopwords
# ----------------------------

generic_resume_words = [
    'date','birth','nationality','indian','place',
    'design','development','role','like','technology',
    'service','environment','experience','summary',
    'projects','project','team','teams','support',
    'skill','responsible','knowledge','tables','reports'
]

custom_stopwords = [
    'people','process','processes','tools','servers',
    'data','table','function','stored','business',
    'custom','report','integration','core','connector',
    'js','web','cs','good','worked','system',
    'team','workded','created','working','involved',
    'support','skill','skills','experience','project',
    'projects','handling','management','knowledge',
    'responsible','server','application','applications',
    'technologies','application server'
]

stop_words = list(
    set(stopwords.words('english'))
    .union(generic_resume_words)
    .union(custom_stopwords)
)

# ----------------------------
# Resume Cleaning Function
# ----------------------------

def preprocess_resume(text):

    if not text:
        return ""

    text = str(text)

    text = text.replace("\\", " ")
    text = re.sub(r'/', ' ', text)
    text = re.sub(r'(?<=\w)[-](?=\w)', '_', text)
    text = re.sub(r'[\)\}\]\(\[\:;,"\'`~!@\$%\^&\*\=<>\?]', ' ', text)

    text = re.sub(r'<.*?>', ' ', text)
    text = re.sub(r'http\S+|www\.\S+', ' ', text)

    text = re.sub(r'\b\d{1,2}(st|nd|rd|th)\b', ' ', text, flags=re.I)
    text = re.sub(r'\b(19|20)\d{2}\b', ' ', text)

    text = re.sub(
        r'\b(jan|january|feb|february|mar|march|apr|april|may|jun|june|'
        r'jul|july|aug|august|sep|sept|september|oct|october|nov|november|dec|december)\b',
        ' ',
        text,
        flags=re.I
    )

    text = re.sub(r'(?<![a-zA-Z])\d+\.\d+(?![a-zA-Z])', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'\b([a-zA-Z]+)[0-9].*?\b', r'\1', text)

    doc = nlp(text)
    clean_tokens = []

    for token in doc:

        t = token.text.lower().strip(".-_")

        if not t:
            continue

        if token.pos_ in {"PRON", "DET", "CCONJ", "SCONJ", "PART", "INTJ"}:
            continue

        if token.ent_type_ in {"DATE", "TIME", "CARDINAL", "QUANTITY", "PERCENT"}:
            continue

        if token.ent_type_ in {"GPE", "LOC", "FAC", "NORP"}:
            continue

        if t in stop_words:
            continue

        if re.fullmatch(r'\d+', t) or re.fullmatch(r'\d+\.\d+', t):
            continue

        if len(t) == 1 and t != 'c':
            continue

        if FLUFF_PATTERN.match(t):
            continue

        if ACTION_VERB_PATTERN.match(t):
            continue

        if TECH_PATTERN.fullmatch(t):
            clean_tokens.append(t)

    return " ".join(clean_tokens)

# ----------------------------
# Experience Extraction
# ----------------------------

import re
from datetime import datetime

def calculate_experience(text):

    text = str(text)

    # -----------------------------------------
    # Direct Years Mention (8 years, 8+ years, 2.5 years)
    # -----------------------------------------
    year_match = re.search(
        r'(\d+(?:\.\d+)?)\s*(\+)?\s*(years?|yrs?)',
        text,
        re.I
    )

    if year_match:
        number = year_match.group(1)
        plus = "+" if year_match.group(2) else ""
        return f"{number}{plus} years"

    # -----------------------------------------
    # Direct Months Mention (9 months)
    # -----------------------------------------
    month_match = re.search(
        r'(\d+)\s*(months?|mos?)',
        text,
        re.I
    )

    if month_match:
        return f"{month_match.group(1)} months"

    # -----------------------------------------
    # Date Range Calculation (MM/YYYY - MM/YYYY or Current)
    # -----------------------------------------
    date_ranges = re.findall(
        r'(\d{2}/\d{4})\s*-\s*(current|present|\d{2}/\d{4})',
        text,
        re.I
    )

    total_months = 0
    now = datetime.now()

    for start_str, end_str in date_ranges:
        try:
            start_date = datetime.strptime(start_str, "%m/%Y")

            if end_str.lower() in ["current", "present"]:
                end_date = now
            else:
                end_date = datetime.strptime(end_str, "%m/%Y")

            months = (end_date.year - start_date.year) * 12 + \
                     (end_date.month - start_date.month)

            if months > 0:
                total_months += months

        except:
            continue

    if total_months > 0:
        years = total_months // 12
        months = total_months % 12

        if years > 0:
            return f"{years} years"
        else:
            return f"{months} months"

    return "Not Found"


# ----------------------------
# Skills Extraction
# ----------------------------

skills_list = [

    # =========================
    # Programming Languages
    # =========================
    "java","c","c++","c#","python","javascript","typescript","shell scripting","powershell",

    # =========================
    # Frontend
    # =========================
    "html","html5","css","css3","bootstrap","react","redux","angular","jquery","material ui",
    "next js","vue","axios","npm",

    # =========================
    # Backend
    # =========================
    "node js","express js",

    # =========================
    # Databases
    # =========================
    "sql","mysql","postgresql","oracle","oracle 10g","oracle 11g","oracle 12c","db2",
    "pl sql","pl/sql","toad","sql developer","oracle apps","r12"

    # =========================
    # PeopleSoft
    # =========================
    "peoplesoft","peoplesoft hrms","peoplesoft fscm","peopletools","application designer","application engine","component interface","integration broker","process scheduler","sqr","cobol","change assistant",
    "file layout","ae","ci","ps query","peoplesoft security","row level security","application package",
    "app engine","peoplecode","nvision","data mover","pia", "ib",                       

    # =========================
    # Workday
    # =========================
    "workday","workday hcm","workday fscm","workday studio","workday integrations","eib","core hcm",
    "calculated fields","report writer", "workday reports","birt","peci","ccw","workday prism","tenant setup"

    # =========================
    # Integration / APIs
    # =========================
    "rest api","rest apis","soap","web services","xml","xslt","json",

    # =========================
    # DevOps / CI-CD
    # =========================
    "git","github","jenkins","ansible","docker","kubernetes","ci cd","bitbucket","maven","gradle","jira","confluence"

    # =========================
    # Middleware / Servers
    # =========================
    "weblogic","tuxedo","apache",

    # =========================
    # Operating Systems
    # =========================
    "linux","unix","windows","aix","oel",

    # =========================
    # Tools / Utilities
    # =========================
    "winscp","filezilla","pcomm","tws","service now","ms office",

    # =========================
    # Cloud
    # =========================
    "aws","ec2","s3","azure","gcp"
]

def extract_skills(text):

    text = text.lower()
    found_skills = []

    for skill in skills_list:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text):
            found_skills.append(skill)

    return sorted(list(set(found_skills)))

# -----------------------------------------------------
# PREDICTION FUNCTION
# -----------------------------------------------------

def predict_resume(raw_text):
    clean_text = preprocess_resume(raw_text)
    encoded_label = model.predict([clean_text])[0]
    predicted_category = label_encoder.inverse_transform([encoded_label])[0]
    return predicted_category

# -----------------------------------------------------
# MULTIPLE FILE UPLOAD
# -----------------------------------------------------

uploaded_files = st.file_uploader(
    "Upload Resumes (PDF, or DOCX)",
    type=["pdf", "docx"],
    accept_multiple_files=True
)

if uploaded_files:

    results = []
    progress_bar = st.progress(0)
    total_files = len(uploaded_files)

    for idx, uploaded_file in enumerate(uploaded_files):

        raw_text = extract_text(uploaded_file)
        if raw_text:
            category = predict_resume(raw_text)
            experience = calculate_experience(raw_text)
            skills = extract_skills(raw_text)

            results.append({
                "File Name": uploaded_file.name,
                "Job Category": category,
                "Experience": experience,
                "Skills": ", ".join(skills)
            })

        progress_bar.progress((idx + 1) / total_files)

    if results:
        st.subheader("📊 Classification Results")
        df_results = pd.DataFrame(results)
        st.dataframe(df_results, use_container_width=True)

        # Save to local folder
        save_path = r"data/resume_classification_results.csv"
        df_results.to_csv(save_path, index=False)

        # Then provide download
        csv = df_results.to_csv(index=False).encode("utf-8")

        st.download_button(
            label="📥 Download Results as CSV",
            data=csv,
            file_name="resume_classification_results.csv",
            mime="text/csv"
        )